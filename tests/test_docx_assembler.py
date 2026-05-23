"""Tests for the DOCX assembler."""
from __future__ import annotations

import pytest
from docx import Document

from writing_suite.agents.docx_assembler import (
    AssemblyOptions,
    DocxAssemblerAgent,
    REF_HEADINGS,
)
from writing_suite.models import (
    Brief,
    Citation,
    Draft,
    DraftSection,
    TokenLedger,
)


@pytest.mark.asyncio
async def test_assembler_creates_docx(tmp_path, sample_brief, sample_draft):
    output = tmp_path / "essay.docx"
    citations = [
        Citation(
            key="Smith_2021",
            authors=["Smith, J."],
            year=2021,
            title="X",
            formatted="Smith, J. (2021). X.",
        )
    ]
    options = AssemblyOptions(output_path=output)

    agent = DocxAssemblerAgent()
    result = await agent.run((sample_brief, sample_draft, citations, options))

    assert result == output
    assert output.exists()
    doc = Document(str(output))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "AI in Modern Education" in text or sample_draft.title in text


@pytest.mark.asyncio
async def test_assembler_uses_language_specific_references_heading(
    tmp_path, sample_draft
):
    output = tmp_path / "id.docx"
    brief_id = Brief(title="T", word_count=1000, language="id")
    citations = [
        Citation(
            key="A_2020",
            authors=["A"],
            year=2020,
            title="x",
            formatted="A (2020). x.",
        )
    ]
    options = AssemblyOptions(output_path=output)

    agent = DocxAssemblerAgent()
    await agent.run((brief_id, sample_draft, citations, options))

    doc = Document(str(output))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert REF_HEADINGS["id"] in text  # Daftar Pustaka


@pytest.mark.asyncio
async def test_assembler_skips_references_when_no_citations(
    tmp_path, sample_brief, sample_draft
):
    output = tmp_path / "noref.docx"
    options = AssemblyOptions(output_path=output)
    agent = DocxAssemblerAgent()
    await agent.run((sample_brief, sample_draft, [], options))

    doc = Document(str(output))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "References" not in text


def test_ref_headings_cover_all_languages():
    """All supported languages should have a localized references heading."""
    for lang in ["en", "id", "ms", "zh"]:
        assert lang in REF_HEADINGS
