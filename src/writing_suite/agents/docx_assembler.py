"""Agent 7: DOCX Assembler.

Pure local logic — no MiMo calls. Takes the final Draft + Citations and
emits a properly formatted .docx file using python-docx. Handles language
direction, citation style headings, and university-style title pages.
"""
from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..models import Brief, Citation, Draft
from ..mimo_client import MimoClient
from ..models import TokenLedger


log = logging.getLogger(__name__)


@dataclass
class AssemblyOptions:
    output_path: Path
    include_title_page: bool = True
    references_heading: str | None = None  # auto-pick by language


REF_HEADINGS = {
    "en": "References",
    "id": "Daftar Pustaka",
    "ms": "Rujukan",
    "zh": "参考文献",
}


class DocxAssemblerAgent:
    """Not a MiMo agent — local docx assembly. Kept in agents/ for symmetry."""

    name = "docx_assembler"

    def __init__(
        self,
        client: MimoClient | None = None,
        ledger: TokenLedger | None = None,
    ):
        # Accepts optional client/ledger to keep the same constructor signature
        # as the MiMo-backed agents, even though they're unused here.
        self.client = client
        self.ledger = ledger

    async def run(
        self, input_data: tuple[Brief, Draft, list[Citation], AssemblyOptions]
    ) -> Path:
        brief, draft, citations, options = input_data
        return self._build(brief, draft, citations, options)

    def _build(
        self,
        brief: Brief,
        draft: Draft,
        citations: list[Citation],
        options: AssemblyOptions,
    ) -> Path:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        if options.include_title_page:
            self._add_title_page(doc, brief, draft)

        for section in draft.sections:
            self._add_section(doc, section)

        if citations:
            self._add_references(doc, citations, brief, options)

        options.output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(options.output_path))
        log.info(
            "wrote docx=%s sections=%d words=%d citations=%d",
            options.output_path,
            len(draft.sections),
            draft.total_words,
            len(citations),
        )
        return options.output_path

    @staticmethod
    def _add_title_page(doc: Document, brief: Brief, draft: Draft) -> None:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(draft.title)
        run.bold = True
        run.font.size = Pt(18)

        if brief.course:
            sub = doc.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub.add_run(brief.course).italic = True

        doc.add_page_break()

    @staticmethod
    def _add_section(doc: Document, section) -> None:
        h = doc.add_heading(section.heading, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for paragraph_text in section.content.split("\n\n"):
            text = paragraph_text.strip()
            if text:
                p = doc.add_paragraph(text)
                p.paragraph_format.first_line_indent = Pt(18)

    @staticmethod
    def _add_references(
        doc: Document,
        citations: list[Citation],
        brief: Brief,
        options: AssemblyOptions,
    ) -> None:
        heading_text = options.references_heading or REF_HEADINGS.get(
            brief.language, "References"
        )
        doc.add_page_break()
        doc.add_heading(heading_text, level=1)
        for cit in sorted(citations, key=lambda c: c.authors[0] if c.authors else c.key):
            text = cit.formatted or f"{', '.join(cit.authors)} ({cit.year}). {cit.title}."
            p = doc.add_paragraph(text)
            p.paragraph_format.first_line_indent = Pt(-18)
            p.paragraph_format.left_indent = Pt(18)
